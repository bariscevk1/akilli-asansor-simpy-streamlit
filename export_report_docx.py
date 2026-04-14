from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "RAPOR.md"
DOCX_PATH = ROOT / "RAPOR.docx"


def _add_runs_with_basic_md(paragraph, text: str) -> None:
    """
    Very small subset:
    - **bold**
    - `inline code`
    """
    token_re = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def build_docx(md_text: str) -> Document:
    doc = Document()

    in_code_block = False
    code_lines: list[str] = []

    for raw in md_text.splitlines():
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # flush code block
                in_code_block = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        img_match = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
        if img_match:
            alt, rel = img_match.group(1), img_match.group(2)
            img_path = (ROOT / rel).resolve()
            if img_path.exists():
                doc.add_paragraph(alt)
                doc.add_picture(str(img_path), width=Inches(2.2))
            else:
                doc.add_paragraph(f"{alt} (görsel bulunamadı: {rel})")
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_basic_md(p, line[2:].strip())
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()
        _add_runs_with_basic_md(p, line.strip())

    return doc


def main() -> None:
    md = MD_PATH.read_text(encoding="utf-8")
    doc = build_docx(md)
    doc.save(str(DOCX_PATH))
    print(f"OK: {DOCX_PATH}")


if __name__ == "__main__":
    main()

